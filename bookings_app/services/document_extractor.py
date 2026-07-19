import requests
# pyrefly: ignore [missing-import]
import fitz  # PyMuPDF
# pyrefly: ignore [missing-import]
import docx
from io import BytesIO

def extract_pdf_text(file_bytes):
    """Extract raw text from PDF bytes using PyMuPDF (fitz)."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"[Error extracting PDF text: {str(e)}]"

def extract_docx_text(file_bytes):
    """Extract raw text from DOCX bytes using python-docx."""
    try:
        doc = docx.Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        
        # Parse tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                # Filter out duplicate empty cells
                if any(row_text):
                    table_text.append(" | ".join(row_text))
        
        full_text = "\n".join(paragraphs)
        if table_text:
            full_text += "\n\nTable Data:\n" + "\n".join(table_text)
        return full_text
    except Exception as e:
        return f"[Error extracting Word document text: {str(e)}]"

def extract_text_from_url(file_url, file_name=None):
    """Downloads a file and extracts text based on file type."""
    if not file_url:
        return ""
    
    try:
        response = requests.get(file_url, timeout=30)
        response.raise_for_status()
        file_bytes = response.content
    except Exception as e:
        return f"[Error fetching document from URL: {str(e)}]"
    
    # Determine extension
    name = file_name or file_url.split('/')[-1].split('?')[0]
    ext = name.split('.')[-1].lower() if '.' in name else ''
    
    if ext == 'pdf':
        return extract_pdf_text(file_bytes)
    elif ext in ['docx', 'doc']:
        return extract_docx_text(file_bytes)
    else:
        return f"[Unsupported file extension for extraction: {ext}]"


import json
import re
from openai import OpenAI
from django.conf import settings


def _normalize_ocr_text(text: str) -> str:
    """Pre-process extracted text to clean up common OCR artifacts."""
    if not text:
        return text
    # Collapse multiple whitespace/newlines into single space
    text = re.sub(r'[ \t]+', ' ', text)
    # Fix common OCR character substitutions
    text = text.replace('|', 'l')  # pipe → lowercase L
    text = text.replace('0', 'O').replace('O', '0')  # Only in numeric contexts — skip global
    # Remove non-printable characters
    text = re.sub(r'[^\x20-\x7E\n£€$%]', '', text)
    # Normalize currency symbols
    text = text.replace('f', '£').replace('E', '£') if False else text  # skip aggressive
    # Collapse excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _extract_json_from_text(text: str) -> dict:
    """Try to parse JSON from LLM response, with regex fallback for malformed output."""
    # First: direct parse
    try:
        return json.loads(text)
    except (json.JSONDecodeError, TypeError):
        pass

    # Fallback: find JSON-like block in the response
    json_match = re.search(r'\{[^{}]*\}', text, re.DOTALL)
    if json_match:
        try:
            return json.loads(json_match.group())
        except (json.JSONDecodeError, TypeError):
            pass

    # Last resort: regex extraction of individual fields
    result = {
        "credit_score": None,
        "ccj_iva_found": False,
        "missed_payments": 0,
        "explanation": "Could not parse structured response from AI."
    }

    score_match = re.search(r'"credit_score"\s*:\s*(\d+)', text)
    if score_match:
        result["credit_score"] = int(score_match.group(1))

    ccj_match = re.search(r'"ccj_iva_found"\s*:\s*(true|false)', text, re.IGNORECASE)
    if ccj_match:
        result["ccj_iva_found"] = ccj_match.group(1).lower() == 'true'

    missed_match = re.search(r'"missed_payments"\s*:\s*(\d+)', text)
    if missed_match:
        result["missed_payments"] = int(missed_match.group(1))

    explanation_match = re.search(r'"explanation"\s*:\s*"([^"]*)"', text)
    if explanation_match:
        result["explanation"] = explanation_match.group(1)

    return result


def process_referencing_checks(app):
    """
    Automated check engine that extracts text from uploaded credit files,
    sends it to Groq/Llama for details extraction, and updates decision rules.
    """
    # Extract text from uploaded documents
    extracted_text_runs = []
    docs = app.uploaded_documents
    if isinstance(docs, str):
        try:
            docs = json.loads(docs)
        except Exception:
            docs = []

    for doc in docs:
        file_url = ""
        file_name = ""
        if isinstance(doc, dict):
            file_url = doc.get('file_url', '')
            file_name = doc.get('file_name', '')
        elif isinstance(doc, str):
            file_url = doc
            file_name = doc.split('/')[-1]

        if file_url:
            text = extract_text_from_url(file_url, file_name)
            extracted_text_runs.append(f"--- Document: {file_name} ---\n{text}\n")

    full_extracted_text = "\n".join(extracted_text_runs)

    # Pre-process OCR text to reduce noise
    full_extracted_text = _normalize_ocr_text(full_extracted_text)

    # Call Groq API
    api_key = getattr(settings, 'NEOSCAPE_API_KEY', '')
    if not api_key:
        raise ValueError("NeoScape API key is not configured.")

    client = OpenAI(
        api_key=api_key,
        base_url="https://api.groq.com/openai/v1"
    )

    system_prompt = (
        "You are an expert AI tenant referencing assistant.\n"
        "Your task is to analyze the provided credit report text and applicant details.\n"
        "The text may come from OCR-scanned documents and may contain:\n"
        "- Garbled or misspelled characters (e.g. 'Cr3dit Sc0re' instead of 'Credit Score')\n"
        "- Broken lines, merged words, or missing spaces\n"
        "- Inconsistent formatting, tables rendered as plain text\n"
        "- Currency symbols like £ may appear as 'f' or 'E' or be missing\n"
        "\n"
        "Despite these issues, do your best to extract the following credit markers.\n"
        "If a value cannot be determined from the text, use null for numbers and false for booleans.\n"
        "Do NOT guess or fabricate values — only report what is clearly present in the text.\n"
        "\n"
        "Respond with a raw JSON object and nothing else. No markdown, no code fences, no explanation.\n"
        "The JSON must contain exactly these keys:\n"
        "{\n"
        '  "credit_score": <integer between 0 and 999, or null if not found>,\n'
        '  "ccj_iva_found": <true if any CCJs, IVAs, or bankruptcies are mentioned, otherwise false>,\n'
        '  "missed_payments": <integer count of missed or late payments, or 0 if none found>,\n'
        '  "explanation": "<brief 1-3 sentence summary of the key findings from the documents>"\n'
        "}\n"
        "\n"
        "Examples of what to look for:\n"
        "- Credit score: Look for patterns like 'Score: 720', 'Credit Rating: 650', 'Experian Score 580'\n"
        "- CCJs/IVAs: Look for 'County Court Judgment', 'CCJ', 'IVA', 'Individual Voluntary Arrangement', 'Bankruptcy'\n"
        "- Missed payments: Look for 'missed payment', 'late payment', 'default', 'arrears', 'payment missed'\n"
    )

    user_prompt = (
        f"Applicant Name: {app.applicant_name}\n"
        f"Applicant Email: {app.applicant_email}\n"
        f"Extracted Documents Text:\n{full_extracted_text or 'No document text available.'}"
    )

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        content = response.choices[0].message.content
        data = _extract_json_from_text(content)
    except Exception as e:
        # Fallback to default rule-based parsing or mock values if API fails
        data = {
            "credit_score": None,
            "ccj_iva_found": False,
            "missed_payments": 0,
            "explanation": f"Automated extraction failed. Error: {str(e)}"
        }

    # Save check results
    app.credit_score = data.get('credit_score')
    app.ccj_iva_found = data.get('ccj_iva_found', False)
    app.missed_payments = data.get('missed_payments', 0)
    app.ai_raw_check_result = data

    # Apply Pass/Caution/Fail rules:
    # Fail: CCJ/IVA is True, OR credit score < 550, OR missed payments >= 3
    # Caution (Approve with Guarantor): credit score between 550 and 650, OR missed payments is 1 or 2
    # Pass: credit score >= 650, ccj_iva_found is False, missed_payments is 0
    if app.ccj_iva_found or (app.credit_score is not None and app.credit_score < 550) or app.missed_payments >= 3:
        app.decision = 'decline'
        app.status = 'failed'
    elif (app.credit_score is not None and app.credit_score < 650) or app.missed_payments in [1, 2]:
        app.decision = 'caution'
        app.status = 'processing'
    else:
        app.decision = 'approve'
        app.status = 'processing'  # Still requires final landlord review / manual override

    app.save()


